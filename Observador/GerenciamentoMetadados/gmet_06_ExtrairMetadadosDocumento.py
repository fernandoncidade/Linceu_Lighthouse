import os
from .gmet_16_ExtrairMetadadosOlefile import extrair_metadados_olefile
from utils.LogManager import LogManager
logger = LogManager.get_logger()

def extrair_metadados_documento(caminho, loc):
    metadados = {}
    ext = os.path.splitext(caminho)[1].lower()

    try:
        if ext == '.pdf':
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(caminho)
                paginas = len(reader.pages)
                metadados['paginas'] = str(paginas)
                linhas_estimadas = paginas * 40
                metadados['linhas'] = str(linhas_estimadas)

                try:
                    linhas_reais = 0
                    palavras_totais = 0
                    for i in range(min(5, paginas)):
                        texto = reader.pages[i].extract_text()
                        linhas_pagina = texto.count('\n') + 1
                        palavras_pagina = len(texto.split())
                        linhas_reais += linhas_pagina
                        palavras_totais += palavras_pagina

                    if paginas > 5:
                        linhas_estimadas = int((linhas_reais / 5) * paginas)
                        palavras_estimadas = int((palavras_totais / 5) * paginas)
                        metadados['linhas'] = str(linhas_estimadas)
                        metadados['palavras'] = str(palavras_estimadas)

                    else:
                        metadados['palavras'] = str(palavras_totais)

                except Exception as e:
                    logger.error(f"Erro ao estimar linhas/palavras do PDF {caminho}: {e}", exc_info=True)

                info = reader.metadata

                if info:
                    if info.title:
                        metadados['titulo'] = info.title

                    if info.author:
                        metadados['autor'] = info.author

                    if info.creator:
                        metadados['criador'] = info.creator

                    if info.producer:
                        metadados['produtor'] = info.producer

                    if info.creation_date:
                        metadados['data_criacao_doc'] = parse_pdf_date(info.creation_date)

                    if info.modification_date:
                        metadados['data_mod_doc'] = parse_pdf_date(info.modification_date)

            except Exception as e:
                logger.error(f"Erro ao extrair metadados do PDF {caminho}: {e}", exc_info=True)

        elif ext in ['.docx', '.doc']:
            try:
                if ext == '.docx':
                    try:
                        from docx import Document
                        doc = Document(caminho)
                        paragrafos = len(doc.paragraphs)
                        palavras_totais = 0
                        linhas_totais = 0

                        for para in doc.paragraphs:
                            texto = para.text.strip()
                            if texto:
                                palavras_no_paragrafo = len(texto.split())
                                palavras_totais += palavras_no_paragrafo
                                linhas_no_paragrafo = max(1, len(texto) // 80 + (1 if len(texto) % 80 > 0 else 0))
                                linhas_totais += linhas_no_paragrafo

                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    texto = cell.text.strip()
                                    if texto:
                                        palavras_totais += len(texto.split())
                                        linhas_no_cell = max(1, len(texto) // 80 + (1 if len(texto) % 80 > 0 else 0))
                                        linhas_totais += linhas_no_cell

                        paginas_estimadas = max(1, palavras_totais // 300)

                        metadados['palavras'] = str(palavras_totais)
                        metadados['linhas'] = str(linhas_totais)
                        metadados['paginas'] = str(paginas_estimadas)
                        metadados['paginas_estimadas'] = str(paginas_estimadas)
                        metadados['paragrafos'] = str(paragrafos)

                        try:
                            props = doc.core_properties
                            if props:
                                if props.title:
                                    metadados['titulo'] = props.title

                                if props.author:
                                    metadados['autor'] = props.author

                                if props.created:
                                    metadados['data_criacao_doc'] = str(props.created)

                                if props.modified:
                                    metadados['data_mod_doc'] = str(props.modified)

                                if props.revision:
                                    metadados['revisao'] = str(props.revision)

                        except Exception as prop_err:
                            if "document is encrypted" in str(prop_err).lower():
                                metadados['protegido'] = loc.get_text("yes") + " (senha)"

                            else:
                                logger.error(f"Erro ao extrair propriedades de DOCX: {prop_err}", exc_info=True)

                    except ImportError as e:
                        logger.error("Biblioteca python-docx não encontrada. Tentando alternativa...", exc_info=True)
                        try:
                            import docx2txt
                            texto_completo = docx2txt.process(caminho)
                            linhas = texto_completo.count('\n') + 1
                            palavras = len(texto_completo.split())
                            paginas_estimadas = max(1, palavras // 300)

                            metadados['palavras'] = str(palavras)
                            metadados['linhas'] = str(linhas)
                            metadados['paginas'] = str(paginas_estimadas)
                            metadados['paginas_estimadas'] = str(paginas_estimadas)

                        except ImportError as e2:
                            logger.error("Biblioteca docx2txt não encontrada. Usando fallback por tamanho.", exc_info=True)
                            tamanho = os.path.getsize(caminho)
                            paginas_estimadas = max(1, tamanho // 20000)
                            palavras_estimadas = paginas_estimadas * 300
                            linhas_estimadas = paginas_estimadas * 40

                            metadados['paginas'] = str(paginas_estimadas)
                            metadados['paginas_estimadas'] = str(paginas_estimadas)
                            metadados['palavras'] = str(palavras_estimadas)
                            metadados['linhas'] = str(linhas_estimadas)

                elif ext == '.doc':
                    try:
                        doc_metadados = extrair_metadados_olefile(caminho, loc)
                        if doc_metadados and 'paginas' in doc_metadados:
                            metadados.update(doc_metadados)

                        else:
                            tamanho = os.path.getsize(caminho)
                            paginas_estimadas = max(1, tamanho // 20000)
                            palavras_estimadas = paginas_estimadas * 300
                            linhas_estimadas = paginas_estimadas * 40

                            metadados['paginas'] = str(paginas_estimadas)
                            metadados['paginas_estimadas'] = str(paginas_estimadas)
                            metadados['palavras'] = str(palavras_estimadas)
                            metadados['linhas'] = str(linhas_estimadas)

                            raise Exception("Fallback para estimativas")

                    except Exception as e:
                        logger.error(f"Usando fallback para olefile: {e}", exc_info=True)

                        import olefile
                        if olefile.isOleFile(caminho):
                            with olefile.OleFileIO(caminho) as ole:
                                tamanho = os.path.getsize(caminho)
                                paginas_estimadas = max(1, tamanho // 20000)
                                palavras_estimadas = paginas_estimadas * 300
                                linhas_estimadas = paginas_estimadas * 40

                                metadados['paginas'] = str(paginas_estimadas)
                                metadados['paginas_estimadas'] = str(paginas_estimadas)
                                metadados['palavras'] = str(palavras_estimadas)
                                metadados['linhas'] = str(linhas_estimadas)

                                if ole.exists('\x05SummaryInformation'):
                                    info = ole.getproperties('\x05SummaryInformation')
                                    if 4 in info:
                                        metadados['autor'] = info[4]

                                    if 2 in info:
                                        metadados['titulo'] = info[2]

                                    if 8 in info:
                                        metadados['ultima_mod'] = info[8]

                                    if 12 in info:
                                        metadados['criacao'] = info[12]

                                    if 19 in info:
                                        metadados['protegido'] = loc.get_text("yes") + f" ({info[19]})"

            except Exception as e:
                logger.error(f"Erro ao extrair metadados do documento {caminho}: {e}", exc_info=True)

        elif ext == '.txt':
            try:
                linhas = 0
                palavras = 0
                caracteres = 0

                with open(caminho, 'r', encoding='utf-8', errors='ignore') as f:
                    for linha in f:
                        linhas += 1
                        palavras += len(linha.split())
                        caracteres += len(linha)

                metadados['linhas'] = str(linhas)
                metadados['palavras'] = str(palavras)
                metadados['caracteres'] = str(caracteres)

            except Exception as e:
                logger.error(f"Erro ao extrair metadados do TXT {caminho}: {e}", exc_info=True)

    except Exception as e:
        logger.error(f"Erro geral ao extrair metadados do documento {caminho}: {e}", exc_info=True)

    return metadados

def parse_pdf_date(date_str):
    import datetime
    try:
        if isinstance(date_str, datetime.datetime):
            return date_str.strftime("%Y-%m-%d %H:%M:%S")

        if isinstance(date_str, str) and date_str.startswith('D:'):
            date_str = date_str[2:]

            year = int(date_str[0:4])
            month = int(date_str[4:6])
            day = int(date_str[6:8])

            if len(date_str) >= 14:
                hour = int(date_str[8:10])
                minute = int(date_str[10:12])
                second = int(date_str[12:14])

                dt = datetime.datetime(year, month, day, hour, minute, second)

            else:
                dt = datetime.datetime(year, month, day)

            return dt.strftime("%Y-%m-%d %H:%M:%S")

        return str(date_str)

    except Exception as e:
        logger.error(f"Erro ao parsear data do PDF: {e}", exc_info=True)
        return str(date_str)
